"""
Exportador de claims a documento Word (.docx) con tabla estructurada.

Genera una tabla con el formato:
| ID | Página | Fecha | Estudio | Tipo | Conclusión | Datos (incluye N) | Evidencia / limitaciones |
"""

import json
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx no está instalado. Ejecuta: pip install python-docx")


# Mapeo de clasificaciones a español
CLASSIFICATION_LABELS = {
    "finding": "Hallazgo",
    "hypothesis": "Hipótesis", 
    "methodological_note": "Nota metodológica"
}


def load_ndjson(ndjson_path: Path) -> tuple[Dict, List[Dict], Dict]:
    """
    Carga un archivo NDJSON y separa meta, claims y summary.
    
    Returns:
        tuple: (meta_record, claims_list, summary_record)
    """
    meta = {}
    claims = []
    summary = {}
    
    with open(ndjson_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            record = json.loads(line)
            record_type = record.get("type")
            
            if record_type == "meta":
                meta = record
            elif record_type == "claim":
                claims.append(record)
            elif record_type == "summary":
                summary = record
    
    return meta, claims, summary


def format_evidence_data(claim: Dict) -> str:
    """Formatea los datos de evidencia para la columna 'Datos (incluye N)'"""
    evidence = claim.get("evidence", {})
    parts = []
    
    # N/Base
    n = evidence.get("n")
    if n:
        parts.append(f"N={n}")
    
    # Tipo de datos
    data_type = evidence.get("data_type")
    if data_type and data_type != "unknown":
        type_labels = {
            "quantitative": "Cuantitativo",
            "qualitative": "Cualitativo", 
            "mixed": "Mixto"
        }
        parts.append(type_labels.get(data_type, data_type))
    
    # Base label si existe
    base_label = evidence.get("base_label")
    if base_label:
        parts.append(f"Base: {base_label}")
    
    return "; ".join(parts) if parts else "No especificado"


def format_evidence_limitations(claim: Dict) -> str:
    """Formatea evidencia y limitaciones para la última columna"""
    parts = []
    
    # Ambiguity flags
    flags = claim.get("ambiguity_flags", [])
    if flags:
        flag_labels = {
            "missing_base": "Sin base/N especificado",
            "low_n_referential": "N bajo (solo referencial)",
            "inferred_n": "N inferido",
            "unspecified_method": "Método no especificado"
        }
        flag_texts = [flag_labels.get(f, f) for f in flags]
        parts.append("Flags: " + ", ".join(flag_texts))
    
    # Classification rationale
    rationale = claim.get("classification_rationale")
    if rationale:
        parts.append(f"Razón: {rationale}")
    
    # Relevance score
    relevance = claim.get("relevance_score")
    if relevance is not None:
        parts.append(f"Relevancia: {relevance:.2f}")
    
    # Theme tags
    tags = claim.get("theme_tags", [])
    if tags:
        parts.append(f"Temas: {', '.join(tags)}")
    
    return "\n".join(parts) if parts else "-"


def set_cell_shading(cell, color: str):
    """Aplica color de fondo a una celda"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def export_to_docx(
    ndjson_path: Path,
    output_path: Optional[Path] = None,
    study_name: Optional[str] = None,
    report_date: Optional[str] = None
) -> Path:
    """
    Exporta los claims de un archivo NDJSON a una tabla Word.
    
    Args:
        ndjson_path: Ruta al archivo .ndjson
        output_path: Ruta de salida (opcional, se genera automáticamente)
        study_name: Nombre del estudio (opcional, se toma del meta)
        report_date: Fecha del informe (opcional, se toma del meta)
    
    Returns:
        Path al archivo .docx generado
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx no está instalado. Ejecuta: pip install python-docx")
    
    # Cargar datos
    meta, claims, summary = load_ndjson(ndjson_path)
    
    # Determinar nombres y fechas
    if not study_name:
        study_name = meta.get("study", {}).get("study_name", "Estudio sin nombre")
    if not report_date:
        report_date = meta.get("study", {}).get("report_date") or \
                      meta.get("extraction", {}).get("extraction_date", "")[:10]
    
    # Crear documento
    doc = Document()
    
    # Título
    title = doc.add_heading(f"Inventario de Conclusiones", level=1)
    
    # Subtítulo con info del estudio
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"Estudio: ").bold = True
    subtitle.add_run(study_name)
    subtitle.add_run(f"\nFecha: ").bold = True
    subtitle.add_run(report_date)
    subtitle.add_run(f"\nTotal claims: ").bold = True
    subtitle.add_run(str(len(claims)))
    
    # Resumen de clasificación
    counts = summary.get("counts", {})
    summary_para = doc.add_paragraph()
    summary_para.add_run("Resumen: ").bold = True
    summary_para.add_run(
        f"{counts.get('findings', 0)} hallazgos | "
        f"{counts.get('hypotheses', 0)} hipótesis | "
        f"{counts.get('methodological_notes', 0)} notas metodológicas"
    )
    
    doc.add_paragraph()  # Espacio
    
    # Crear tabla
    # Columnas: ID | Página | Fecha | Estudio | Tipo | Conclusión | Datos | Evidencia/limitaciones
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Encabezados
    headers = ["ID", "Página", "Fecha", "Estudio", "Tipo", "Conclusión", "Datos (N)", "Evidencia / limitaciones"]
    header_cells = table.rows[0].cells
    
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        # Formato de encabezado
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        # Color de fondo gris
        set_cell_shading(cell, "D9D9D9")
    
    # Filas de datos
    for claim in claims:
        row = table.add_row()
        cells = row.cells
        
        # ID
        cells[0].text = claim.get("id", "-")
        
        # Página
        page = claim.get("page_number")
        cells[1].text = f"p.{page}" if page else "-"
        
        # Fecha (del estudio)
        cells[2].text = report_date
        
        # Estudio
        cells[3].text = study_name[:50] + "..." if len(study_name) > 50 else study_name
        
        # Tipo (con negrita)
        classification = claim.get("classification", "hypothesis")
        tipo_label = CLASSIFICATION_LABELS.get(classification, classification)
        cells[4].text = tipo_label
        
        # Aplicar negrita al texto de tipo
        for paragraph in cells[4].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
        
        # Color según tipo
        if classification == "finding":
            set_cell_shading(cells[4], "C6EFCE")  # Verde claro
        elif classification == "hypothesis":
            set_cell_shading(cells[4], "FFEB9C")  # Amarillo claro
        else:
            set_cell_shading(cells[4], "DDEBF7")  # Azul claro
        
        # Conclusión
        cells[5].text = claim.get("claim_text", "-")
        
        # Datos (incluye N)
        cells[6].text = format_evidence_data(claim)
        
        # Evidencia / limitaciones
        cells[7].text = format_evidence_limitations(claim)
        
        # Formato de celdas (excepto columna Tipo que ya tiene formato)
        for i, cell in enumerate(cells):
            if i == 4:  # Saltar columna Tipo, ya tiene formato
                continue
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    # Ajustar anchos de columna (aproximados)
    # Total ~18cm para A4 con márgenes
    widths = [Cm(1.2), Cm(1.2), Cm(2.0), Cm(3.0), Cm(1.8), Cm(4.5), Cm(2.3), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    
    # Generar ruta de salida
    if output_path is None:
        output_path = ndjson_path.with_suffix('.docx')
    
    # Guardar
    doc.save(output_path)
    
    return output_path


def export_from_analysis(
    ndjson_path: str,
    output_path: Optional[str] = None,
    study_name: Optional[str] = None,
    report_date: Optional[str] = None
) -> str:
    """
    Wrapper para exportar desde línea de comandos o main.py
    
    Args:
        ndjson_path: Ruta al archivo .ndjson (string)
        output_path: Ruta de salida opcional
        study_name: Nombre del estudio
        report_date: Fecha del informe
    
    Returns:
        Ruta al archivo .docx generado (string)
    """
    ndjson_p = Path(ndjson_path)
    output_p = Path(output_path) if output_path else None
    
    result = export_to_docx(ndjson_p, output_p, study_name, report_date)
    return str(result)


# CLI standalone
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Exporta claims de NDJSON a tabla Word (.docx)"
    )
    parser.add_argument(
        "ndjson_file",
        help="Ruta al archivo .ndjson generado por main.py"
    )
    parser.add_argument(
        "-o", "--output",
        help="Ruta de salida para el .docx (opcional)"
    )
    parser.add_argument(
        "--study-name",
        help="Nombre del estudio (opcional, se toma del archivo)"
    )
    parser.add_argument(
        "--report-date",
        help="Fecha del informe (opcional, se toma del archivo)"
    )
    
    args = parser.parse_args()
    
    try:
        output = export_from_analysis(
            args.ndjson_file,
            args.output,
            args.study_name,
            args.report_date
        )
        print(f"✅ Exportado: {output}")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
